"""
文档内容解析服务 —— 支持 PDF / Word / TXT 格式
"""

import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


def parse_pdf(data: bytes) -> str:
    """解析 PDF 文件，返回纯文本"""
    from PyPDF2 import PdfReader

    reader = PdfReader(io.BytesIO(data))
    total_pages = len(reader.pages)
    logger.info(f"[解析] PDF 共 {total_pages} 页，开始逐页提取文本...")
    pages_text = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            pages_text.append(text.strip())
        if (i + 1) % 10 == 0 or (i + 1) == total_pages:
            logger.info(f"[解析] PDF 页面进度: {i+1}/{total_pages}")
    full_text = "\n\n".join(pages_text)
    logger.info(f"[解析] PDF 解析完成，有效页 {len(pages_text)}/{total_pages}，文本 {len(full_text)} 字符")
    return full_text


def parse_docx(data: bytes) -> str:
    """解析 Word (.docx) 文件，返回纯文本"""
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(data))
    logger.info(f"[解析] DOCX 段落数: {len(doc.paragraphs)}, 表格数: {len(doc.tables)}")
    paragraphs_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs_text.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs_text.append(row_text)

    full_text = "\n".join(paragraphs_text)
    logger.info(f"[解析] DOCX 解析完成，有效段落 {len(paragraphs_text)} 个，文本 {len(full_text)} 字符")
    return full_text


def parse_txt(data: bytes) -> str:
    """解析纯文本文件"""
    # 尝试 UTF-8，失败则回退到 GBK
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            text = data.decode(encoding)
            logger.info(f"TXT 解析完成（编码: {encoding}），提取文本 {len(text)} 字符")
            return text
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError("无法识别文件编码")


# 格式 -> 解析函数 映射
_PARSERS = {
    "pdf": parse_pdf,
    "docx": parse_docx,
    "doc": parse_docx,  # .doc 也尝试用 docx 解析
    "txt": parse_txt,
}


def parse_document(data: bytes, file_format: str) -> str:
    """
    根据文件格式解析文档内容

    Args:
        data: 文件二进制内容
        file_format: 文件格式（pdf / docx / txt）

    Returns:
        提取的纯文本内容

    Raises:
        ValueError: 不支持的文件格式
    """
    fmt = file_format.lower().strip(".")
    parser = _PARSERS.get(fmt)
    if parser is None:
        raise ValueError(f"不支持的文件格式: {file_format}，当前支持: {list(_PARSERS.keys())}")
    return parser(data)
